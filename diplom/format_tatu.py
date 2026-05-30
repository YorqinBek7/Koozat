"""TATU diplom ishi standartiga muvofiq .docx formatlash."""
import sys
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def apply_tatu_format(input_path: str, output_path: str) -> None:
    doc = Document(input_path)

    # Sahifa parametrlari: A4, TATU standart maydalar
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)

    # "Normal" uslubni Times New Roman 14pt, 1.5 interval qilib o'rnatish
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(14)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # Har bir paragrafni formatlash
    for para in doc.paragraphs:
        style_name = para.style.name
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            rPr2 = run._element.get_or_add_rPr()
            rFonts2 = rPr2.find(qn('w:rFonts'))
            if rFonts2 is None:
                rFonts2 = OxmlElement('w:rFonts')
                rPr2.insert(0, rFonts2)
            rFonts2.set(qn('w:ascii'), 'Times New Roman')
            rFonts2.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts2.set(qn('w:cs'), 'Times New Roman')

        pfp = para.paragraph_format
        pfp.line_spacing = 1.5
        pfp.space_before = Pt(0)
        pfp.space_after = Pt(0)

        if style_name == 'Normal' and para.text.strip():
            pfp.first_line_indent = Cm(1.25)
            pfp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif 'Heading 1' in style_name:
            pfp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pfp.space_before = Pt(12)
            pfp.space_after = Pt(12)
            for run in para.runs:
                run.font.size = Pt(16)
                run.font.bold = True
        elif 'Heading 2' in style_name:
            pfp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pfp.space_before = Pt(10)
            pfp.space_after = Pt(6)
            for run in para.runs:
                run.font.size = Pt(15)
                run.font.bold = True
        elif 'Heading 3' in style_name:
            pfp.space_before = Pt(8)
            pfp.space_after = Pt(4)
            for run in para.runs:
                run.font.size = Pt(14)
                run.font.bold = True

    # Jadvallarga ramka qo'shish va shrift formatlash
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_cell_borders(cell)
                for para in cell.paragraphs:
                    pfp = para.paragraph_format
                    pfp.line_spacing = 1.15
                    pfp.first_line_indent = Cm(0)
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    doc.save(output_path)
    print(f"Saqlandi: {output_path}")


if __name__ == '__main__':
    inp = sys.argv[1] if len(sys.argv) > 1 else '04_1_bob.docx'
    outp = sys.argv[2] if len(sys.argv) > 2 else '04_1_bob_tatu.docx'
    apply_tatu_format(inp, outp)
