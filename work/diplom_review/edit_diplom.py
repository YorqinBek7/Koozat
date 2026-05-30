from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


SOURCE = Path("/Users/yorqin/Desktop/DIPLOM_Yuldashev_Yorqin.docx")
OUTPUT = Path("/Users/yorqin/Documents/Projects/koozat/work/diplom_review/DIPLOM_Yuldashev_Yorqin_tahrirlangan.docx")


REPLACEMENTS = {
    "Ekoperimentning": "Eksperimentning",
    "Notog‘ri": "Noto‘g‘ri",
    "notog‘ri": "noto‘g‘ri",
    "soviq ishga tushishi": "dastlabki ishga tushish vaqti",
    "soviq ishga tushishi (cold start)": "dastlabki ishga tushish vaqti (cold start)",
    "roditellari": "ota-onalari",
    "roditellar": "ota-onalar",
    "Roditellar": "Ota-onalar",
    "ihtiyoriy": "ixtiyoriy",
    "foydalanuvchi safari": "foydalanuvchi ssenariysi",
    "user journey (foydalanuvchi safari)": "foydalanuvchi ssenariysi (user journey)",
    "implementatsiyasi": "amaliy joriy etilishi",
    "implementatsiyasini": "amaliy joriy etilishini",
    "Repozitoriy implementatsiyasi": "Repozitoriy amaliy joriy etilishi",
    "Performans optimallashtirishlari": "Ishlash samaradorligini optimallashtirish",
    "platforma bilan qanday tezlikda muloqotda bo‘lishini": "platformadan qanchalik tez-tez foydalanishini",
    "24 nafar talabani va 2 nafar o‘qituvchini": "24 nafar talaba va 2 nafar o‘qituvchini",
    "Adabiyot va to‘liq foydalanuvchi tajribasi yordamida": "Adabiyotlar tahlili va foydalanuvchi tajribasi asosida",
    "o‘zaro aloqaviy dizayni": "interaktiv dizayni",
    "digital native avlodga mansub": "raqamli muhitda ulg‘aygan avlodga mansub",
    "cross-platform UI framework": "kross-platformali UI freymvork",
    "cloud backend": "bulutli backend",
    "Flutter framework": "Flutter freymvorki",
    "Veb-platforma qo‘llovi": "Veb-platformani qo‘llab-quvvatlash",
    "mobil-birinchi yondashuv": "mobil qurilmalarga ustuvor yondashuv",
    "self-hosted": "o‘z serverida joylashtiriladigan",
    "Self-hosted": "O‘z serverida joylashtiriladigan",
    "vendor lock-in": "bitta xizmat ko‘rsatuvchiga bog‘lanib qolish",
    "vendor-lock-in": "bitta xizmat ko‘rsatuvchiga bog‘lanib qolish",
    "multitenancy": "ko‘p tashkilotli arxitektura",
    "Multitenancy": "Ko‘p tashkilotli arxitektura",
    "Multitenancy SaaS": "Ko‘p tashkilotli SaaS modeli",
    "AI-prediktiv analitika": "sun’iy intellekt asosidagi prediktiv analitika",
    "AI-prediktiv": "sun’iy intellekt asosidagi prediktiv",
    "gey­mifikatsiya": "gey­mifikatsiya",
    "Geymifikatsiya": "O‘yinlashtirish",
    "geymifikatsiya": "o‘yinlashtirish",
    "adaptiv ta’lim": "moslashuvchan ta’lim",
    "Adaptiv ta’lim": "Moslashuvchan ta’lim",
    "intervensiya": "pedagogik aralashuv",
    "Intervensiya": "Pedagogik aralashuv",
    "operativ tuzatishlar": "tezkor tuzatishlar",
    "operativ ma’lumot": "tezkor ma’lumot",
    "operativlashtirish": "tezlashtirish",
    "sub’ektiv": "subyektiv",
    "ob’ektiv": "obyektiv",
    "ob’yekti": "obyekti",
    "ob’yekt": "obyekt",
    "avtentifikatsiya": "autentifikatsiya",
    "milisekund": "millisekund",
    "Stuydent t-testi": "Student t-testi",
    "to‘rt asosiy ekranlar guruhi": "to‘rt asosiy ekran guruhi",
    "Bu uch ko‘rsatkichlar guruhini": "Bu uch ko‘rsatkich guruhini",
    "Hemis": "HEMIS",
    "Moodleda": "Moodle tizimida",
    "Google Classroomda": "Google Classroom tizimida",
    "foydalanuvchini ro‘yxatdan o‘tkazish": "foydalanuvchini ro‘yxatdan o‘tkazish",
    "Ish joyi (/Users/yorqin/Documents/Projects/koozat/ katalogiga teng masofadagi jismoniy ish stoli) ergonomik talablarga muvofiq sozlangan: ": "Ish joyi ergonomik talablarga muvofiq sozlangan: ",
    "sodda lekin kuchli analitika": "sodda, ammo kuchli analitika",
    "kichik testlash sirti": "cheklangan sinov qamrovi",
    "to‘liq foydalanuvchi tajribasi": "foydalanuvchi tajribasi",
    "Pedagogik pedagogik aralashuv": "Pedagogik aralashuv",
    "pedagogik pedagogik aralashuv": "pedagogik aralashuv",
    "Haqiqiy real vaqtli": "Real vaqtli",
    "haqiqiy real vaqtli": "real vaqtli",
    "Oyoq oyog‘i": "Oyoq kafti",
    "masshtashlash": "masshtablash",
    "juda yuqori darajada ahamiyatli": "yuqori darajada ahamiyatli",
    "juda yuqori ahamiyatli": "yuqori darajada ahamiyatli",
    "juda kech": "kech",
    "juda foydali": "samarali",
    "eng kuchli omillardan biri": "muhim omillardan biri",
    "keskin qisqarib": "sezilarli qisqarib",
    "keskin qisqarishi": "sezilarli qisqarishi",
    "talabalar progressi": "talabalarning o‘quv natijalari",
    "Talaba progressining": "Talaba o‘quv natijalarining",
    "Talaba progressi": "Talaba o‘quv natijalari",
    "talabaning progressini": "talabaning o‘quv natijalarini",
    "Talabaga o‘z progressini": "Talabaga o‘z o‘quv natijalarini",
    "o‘z progressini": "o‘z o‘quv natijalarini",
    "shaxsiy o‘quv progressi": "shaxsiy o‘quv natijalari",
    "ishlangan progress grafiklari": "ishlangan o‘quv natijalari grafiklari",
    "talabalar o‘quv natijalari": "talabalarning o‘quv natijalari",
}

RAW_REPLACEMENTS = {
    "Talaba progressining": "Talaba o‘quv natijalarining",
    "talabalar progressi": "talabalarning o‘quv natijalari",
    "Talabaga o‘z progressini": "Talabaga o‘z o‘quv natijalarini",
}


def iter_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def set_text(paragraph: Paragraph, text: str) -> None:
    paragraph.text = text


def insert_after(paragraph: Paragraph, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        try:
            new_para.style = style
        except Exception:
            pass
    new_para.text = text
    return new_para


def apply_replacements(text: str) -> str:
    updated = text
    for old, new in REPLACEMENTS.items():
        updated = updated.replace(old, new)
    updated = updated.replace(" # IV BOB. HAYOT FAOLIYATI XAVFSIZLIGI", "\nIV BOB. HAYOT FAOLIYATI XAVFSIZLIGI")
    updated = updated.replace(" # FOYDALANILGAN ADABIYOTLAR RO‘YXATI", "\nFOYDALANILGAN ADABIYOTLAR RO‘YXATI")
    return updated


def main() -> None:
    doc = Document(SOURCE)

    for paragraph in list(iter_paragraphs(doc)):
        original = paragraph.text
        if not original:
            continue

        updated = apply_replacements(original)

        if "\nIV BOB. HAYOT FAOLIYATI XAVFSIZLIGI" in updated:
            before, after = updated.split("\nIV BOB. HAYOT FAOLIYATI XAVFSIZLIGI", 1)
            set_text(paragraph, before.strip())
            insert_after(paragraph, "IV BOB. HAYOT FAOLIYATI XAVFSIZLIGI" + after, style=paragraph.style)
            continue

        if "\nFOYDALANILGAN ADABIYOTLAR RO‘YXATI" in updated:
            before, after = updated.split("\nFOYDALANILGAN ADABIYOTLAR RO‘YXATI", 1)
            set_text(paragraph, before.strip())
            insert_after(paragraph, "FOYDALANILGAN ADABIYOTLAR RO‘YXATI" + after, style=paragraph.style)
            continue

        if updated != original:
            set_text(paragraph, updated)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    with ZipFile(OUTPUT, "r") as zin, NamedTemporaryFile(delete=False, suffix=".docx", dir=OUTPUT.parent) as tmp:
        tmp_path = Path(tmp.name)
        with ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    text = data.decode("utf-8")
                    for old, new in RAW_REPLACEMENTS.items():
                        text = text.replace(old, new)
                    data = text.encode("utf-8")
                zout.writestr(item, data)
    tmp_path.replace(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
